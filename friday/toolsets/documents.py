"""
Reading the files `files_read` cannot.

`files_read` handles UTF-8 text and reports everything else as binary, so
"what's in this PDF?" had no answer at all. This is the donor capability worth
taking from `Mark-L/actions/file_processor.py`, reduced to the part that is
deterministic.

What was deliberately left behind, having read the source rather than the
summary of it:

  action="run"   `subprocess.run(["python", str(path)])`. An arbitrary
                 execution primitive reachable by naming a file - the same
                 class as the `{"shell": ...}` step that kept Mark-L's
                 automation engine out. It does not come across, and there is
                 no equivalent here.
  the model calls  the donor summarises, translates and "analyses" by handing
                 file contents to Gemini inside the tool. Friday already has a
                 model, and it is better placed to decide what to do with text
                 than a file handler is. These return the text; the summarising
                 is a turn, not a tool.
  audio, video   both need ffmpeg, which is not installed on this machine -
                 measured, and the reason music cannot play tonight. A
                 capability that cannot run is not worth a routing slot.

Everything here goes through the same filesystem jail as `files_read`: the
path is chosen by an untrusted caller, and a document reader pointed at
`.env` is an exfiltration primitive wearing a helpful name.
"""

from __future__ import annotations

import csv
import io
import json
import zipfile
from pathlib import Path

from friday import contracts as c
from friday.fsjail import JailError
from friday.policy import PolicyEngine, default_engine
from friday.toolsets.files import jail
from friday.toolsets.system import APPROVAL_PREFIX

EXECUTION_SCOPE = "local_machine"

#: Text is returned to a language model through a voice turn. Past a point,
#: more of it is cost rather than answer - and the caller can ask for a
#: specific page or sheet when it needs more.
MAX_TEXT_CHARS = 40_000

#: Rows from a spreadsheet, before it stops being a preview and starts being
#: a dataset. Something with more rows than this wants product processing or
#: a real query, not a tool that reads it into a voice reply.
MAX_ROWS = 200

PDF = ".pdf"
DOCX = ".docx"
XLSX = frozenset({".xlsx", ".xlsm"})
PPTX = ".pptx"
ARCHIVES = frozenset({".zip"})

KNOWN = {PDF, DOCX, PPTX} | XLSX | ARCHIVES


class DocumentError(RuntimeError):
    """The file is not readable as the kind of document it claims to be."""


def _gate(run: c.Run, tool_id: str, engine: PolicyEngine) -> c.ActionResult | None:
    verdict = engine.decide(tool_id)
    if verdict.allowed:
        return None
    return run.record(c.started(run.run_id, tool_id).finish(
        status=c.CANCELLED,
        error=f"{APPROVAL_PREFIX}: {verdict.reason} [{verdict.decision}]",
    ))


def _scoped(payload: dict) -> dict:
    return {"execution_scope": EXECUTION_SCOPE, **payload}


def _safe(run: c.Run, started: c.ActionResult, raw: str):
    try:
        return jail().resolve(raw), None
    except JailError as exc:
        return None, run.record(c.failed(started, f"path refused: {exc}"))


# ---------------------------------------------------------------------------
# Extraction, one function per format
#
# Each returns (text, facts). `facts` is what the format knows about itself -
# page count, sheet names - which is the part a summary cannot reconstruct.
# ---------------------------------------------------------------------------


def _from_pdf(path: Path, pages: str = "") -> tuple[str, dict]:
    from pypdf import PdfReader

    reader = PdfReader(str(path))
    if reader.is_encrypted:
        try:
            reader.decrypt("")
        except Exception as exc:                       # noqa: BLE001
            raise DocumentError(f"{path.name} is encrypted") from exc

    wanted = _page_numbers(pages, len(reader.pages))
    chunks = []
    for number in wanted:
        try:
            chunks.append(reader.pages[number - 1].extract_text() or "")
        except Exception:                              # noqa: BLE001
            chunks.append("")
    text = "\n\n".join(chunks).strip()
    facts = {"pages": len(reader.pages), "pages_read": wanted,
             # A scanned PDF is images of text. Saying "0 characters" is not
             # the same as saying "this one needs OCR, which I do not have".
             "looks_scanned": not text and bool(reader.pages)}
    return text, facts


def _page_numbers(pages: str, total: int) -> list[int]:
    """"3", "2-5", "" -> a list of 1-based page numbers that exist."""
    wanted = (pages or "").strip()
    if not wanted:
        return list(range(1, total + 1))
    numbers: list[int] = []
    for part in wanted.replace(" ", "").split(","):
        if "-" in part:
            start, _, end = part.partition("-")
            if start.isdigit() and end.isdigit():
                numbers.extend(range(int(start), int(end) + 1))
        elif part.isdigit():
            numbers.append(int(part))
    return [n for n in numbers if 1 <= n <= total] or list(range(1, total + 1))


def _from_docx(path: Path) -> tuple[str, dict]:
    import docx

    document = docx.Document(str(path))
    paragraphs = [p.text for p in document.paragraphs if p.text.strip()]
    tables = 0
    for table in document.tables:
        tables += 1
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            if any(cells):
                paragraphs.append(" | ".join(cells))
    return "\n".join(paragraphs).strip(), {
        "paragraphs": len(document.paragraphs), "tables": tables}


def _from_pptx(path: Path) -> tuple[str, dict]:
    """
    Slides, without python-pptx.

    A .pptx is a zip of XML, and the text of a slide is the `<a:t>` elements
    in order. That is a whole dependency avoided for about ten lines - and
    unlike the library it cannot fail to install on a machine that will only
    ever read a handful of decks.
    """
    import re

    text_run = re.compile(r"<a:t>(.*?)</a:t>", re.S)
    unescape = {"&amp;": "&", "&lt;": "<", "&gt;": ">", "&quot;": '"',
                "&apos;": "'"}

    slides: list[str] = []
    with zipfile.ZipFile(path) as archive:
        names = sorted(
            (n for n in archive.namelist()
             if n.startswith("ppt/slides/slide") and n.endswith(".xml")),
            key=lambda n: int("".join(ch for ch in n if ch.isdigit()) or 0))
        for name in names:
            xml = archive.read(name).decode("utf-8", errors="ignore")
            runs = [text_run.sub(lambda m: m.group(1), match)
                    for match in text_run.findall(xml)]
            body = " ".join(runs).strip()
            for entity, character in unescape.items():
                body = body.replace(entity, character)
            slides.append(body)
    text = "\n\n".join(f"[slide {i}] {body}"
                       for i, body in enumerate(slides, 1) if body)
    return text.strip(), {"slides": len(slides)}


def _from_xlsx(path: Path, sheet: str = "") -> tuple[str, dict]:
    from openpyxl import load_workbook

    # read_only because a spreadsheet can be very large and this only reads;
    # data_only because a formula's text is not what anybody asked for.
    workbook = load_workbook(str(path), read_only=True, data_only=True)
    try:
        names = list(workbook.sheetnames)
        wanted = sheet.strip() or names[0]
        if wanted not in names:
            raise DocumentError(
                f"{path.name} has no sheet {sheet!r}; it has {names}")
        worksheet = workbook[wanted]

        buffer = io.StringIO()
        writer = csv.writer(buffer)
        rows = 0
        for row in worksheet.iter_rows(values_only=True):
            if rows >= MAX_ROWS:
                break
            if any(cell is not None and str(cell).strip() for cell in row):
                writer.writerow(["" if cell is None else cell for cell in row])
                rows += 1
        return buffer.getvalue().strip(), {
            "sheets": names, "sheet_read": wanted, "rows_read": rows,
            "truncated_rows": rows >= MAX_ROWS}
    finally:
        workbook.close()


def _from_zip(path: Path) -> tuple[str, dict]:
    """
    What is in the archive, not what is in the files.

    Listing is a read; extracting is a write, and a zip that unpacks
    `../../../.ssh/id_rsa` is the oldest trick there is. Extraction belongs in
    its own tool, through the jail, and does not exist yet.
    """
    with zipfile.ZipFile(path) as archive:
        entries = [{"name": info.filename, "bytes": info.file_size,
                    "compressed": info.compress_size}
                   for info in archive.infolist() if not info.is_dir()]
    listing = "\n".join(f"{e['name']}  ({e['bytes']} bytes)"
                        for e in entries[:MAX_ROWS])
    return listing, {"entries": len(entries),
                     "total_bytes": sum(e["bytes"] for e in entries)}


# ---------------------------------------------------------------------------
# The tool
# ---------------------------------------------------------------------------


def documents_extract(
    run: c.Run, path: str, *, pages: str = "", sheet: str = "",
    max_chars: int = MAX_TEXT_CHARS, engine: PolicyEngine = default_engine,
) -> c.ActionResult:
    """
    Read a PDF, Word document, spreadsheet, slide deck or archive as text.
    """
    tool_id = "documents.extract"
    blocked = _gate(run, tool_id, engine)
    if blocked:
        return blocked
    started = c.started(run.run_id, tool_id)

    target, failure = _safe(run, started, path)
    if failure:
        return failure
    if not target.is_file():
        return run.record(c.failed(started, f"no such file: {target}"))

    suffix = target.suffix.lower()
    if suffix not in KNOWN:
        return run.record(c.failed(
            started,
            f"{target.name} is a {suffix or 'suffixless'} file; this reads "
            f"{sorted(KNOWN)}. Plain text goes through files_read."))

    try:
        if suffix == PDF:
            text, facts = _from_pdf(target, pages)
        elif suffix == DOCX:
            text, facts = _from_docx(target)
        elif suffix == PPTX:
            text, facts = _from_pptx(target)
        elif suffix in XLSX:
            text, facts = _from_xlsx(target, sheet)
        else:
            text, facts = _from_zip(target)
    except DocumentError as exc:
        return run.record(c.failed(started, str(exc)))
    except Exception as exc:                            # noqa: BLE001
        return run.record(c.failed(
            started, f"could not read {target.name} as {suffix}: "
                     f"{type(exc).__name__}: {exc}"))

    if facts.get("looks_scanned"):
        # PARTIAL, not FAILED: the file was read correctly and contains no
        # text layer. Reporting "0 characters" would be true and useless.
        return run.record(c.partial(
            started,
            f"{target.name} has {facts['pages']} page(s) and no text layer - "
            f"it is scanned images, and reading those needs OCR, which Friday "
            f"does not have",
            output=_scoped({"path": str(target), "kind": suffix.lstrip("."),
                            "chars": 0, "text": "", **facts})))

    truncated = len(text) > max_chars
    return run.record(c.succeeded(
        started,
        output=_scoped({"path": str(target), "kind": suffix.lstrip("."),
                        "text": text[:max_chars], "chars": len(text),
                        "truncated": truncated, **facts}),
        verification=c.Verification(
            method="document_extract",
            evidence=f"{target.name}: {len(text)} characters from "
                     f"{suffix.lstrip('.')}"
                     + (f", {facts['pages']} page(s)" if "pages" in facts else "")
                     + (f", sheet {facts['sheet_read']!r} of "
                        f"{len(facts['sheets'])}" if "sheets" in facts else "")
                     + (f", {facts['slides']} slide(s)" if "slides" in facts else "")
                     + (f", {facts['entries']} entr(ies)" if "entries" in facts else "")
                     + ("; truncated" if truncated else ""),
        ),
    ))


def documents_inspect(
    run: c.Run, path: str, *, engine: PolicyEngine = default_engine
) -> c.ActionResult:
    """
    What a document is, without reading all of it.

    Separate from extraction because "how long is this?" and "what does it
    say?" are different questions with very different costs, and answering the
    first by doing the second is how a voice turn spends thirty seconds.
    """
    tool_id = "documents.inspect"
    blocked = _gate(run, tool_id, engine)
    if blocked:
        return blocked
    started = c.started(run.run_id, tool_id)

    target, failure = _safe(run, started, path)
    if failure:
        return failure
    if not target.is_file():
        return run.record(c.failed(started, f"no such file: {target}"))

    suffix = target.suffix.lower()
    if suffix not in KNOWN:
        return run.record(c.failed(
            started, f"{target.name} is not a document this reads; "
                     f"known: {sorted(KNOWN)}"))

    facts: dict = {"kind": suffix.lstrip("."),
                   "size_bytes": target.stat().st_size}
    try:
        if suffix == PDF:
            from pypdf import PdfReader

            reader = PdfReader(str(target))
            facts["pages"] = len(reader.pages)
            facts["encrypted"] = reader.is_encrypted
            information = reader.metadata or {}
            facts["title"] = str(information.get("/Title", "") or "")
        elif suffix == DOCX:
            import docx

            document = docx.Document(str(target))
            facts["paragraphs"] = len(document.paragraphs)
            facts["tables"] = len(document.tables)
        elif suffix == PPTX:
            _, pptx_facts = _from_pptx(target)
            facts.update(pptx_facts)
        elif suffix in XLSX:
            from openpyxl import load_workbook

            workbook = load_workbook(str(target), read_only=True)
            try:
                facts["sheets"] = list(workbook.sheetnames)
            finally:
                workbook.close()
        else:
            with zipfile.ZipFile(target) as archive:
                infos = [i for i in archive.infolist() if not i.is_dir()]
            facts["entries"] = len(infos)
            facts["total_bytes"] = sum(i.file_size for i in infos)
    except Exception as exc:                            # noqa: BLE001
        return run.record(c.failed(
            started, f"could not inspect {target.name}: "
                     f"{type(exc).__name__}: {exc}"))

    return run.record(c.succeeded(
        started,
        output=_scoped({"path": str(target), **facts}),
        verification=c.Verification(
            method="document_inspect",
            evidence=f"{target.name}: "
                     + json.dumps({k: v for k, v in facts.items()
                                   if k != "path"}, default=str)[:200],
        ),
    ))
