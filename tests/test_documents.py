"""
Reading the files `files_read` cannot.

Every document here is built in the test, with the real libraries, and read
back with the real readers. A fixture PDF checked into the repo proves that a
particular file still parses; building one proves the round trip.

What is NOT here is as much the point. The donor this came from
(`Mark-L/actions/file_processor.py`) has `action="run"` ->
`subprocess.run(["python", path])`, an arbitrary execution primitive reachable
by naming a file. It did not come across, and the last test in this file is
what keeps it out.
"""

from __future__ import annotations

import zipfile

import pytest

from friday import contracts as c
from friday import policy as P
from friday.fsjail import FileJail
from friday.toolsets import documents as D
from friday.toolsets import files as F


@pytest.fixture
def root(tmp_path):
    F.reset_jail(FileJail(roots=(tmp_path,)))
    yield tmp_path
    F.reset_jail(None)


@pytest.fixture
def run():
    return c.Run.create("read that document", capability="documents")


# ---------------------------------------------------------------------------
# Builders: real files, made by the real libraries
# ---------------------------------------------------------------------------


def a_pdf(path, pages=("Arc reactor maintenance log", "Palladium core notes")):
    from pypdf import PdfWriter

    # A page with text, written by a real PDF library rather than by hand.
    from pypdf.generic import (ArrayObject, DecodedStreamObject, DictionaryObject,
                               FloatObject, NameObject, NumberObject)

    writer = PdfWriter()
    for body in pages:
        page = writer.add_blank_page(width=200, height=200)
        stream = DecodedStreamObject()
        stream.set_data(
            f"BT /F1 12 Tf 20 100 Td ({body}) Tj ET".encode("latin-1"))
        font = DictionaryObject({
            NameObject("/Type"): NameObject("/Font"),
            NameObject("/Subtype"): NameObject("/Type1"),
            NameObject("/BaseFont"): NameObject("/Helvetica"),
        })
        page[NameObject("/Contents")] = writer._add_object(stream)
        page[NameObject("/Resources")] = DictionaryObject({
            NameObject("/Font"): DictionaryObject({
                NameObject("/F1"): writer._add_object(font)})})
        page[NameObject("/MediaBox")] = ArrayObject(
            [NumberObject(0), NumberObject(0),
             FloatObject(200), FloatObject(200)])
    with path.open("wb") as handle:
        writer.write(handle)
    return path


def a_blank_pdf(path, pages=3):
    from pypdf import PdfWriter

    writer = PdfWriter()
    for _ in range(pages):
        writer.add_blank_page(width=200, height=200)
    with path.open("wb") as handle:
        writer.write(handle)
    return path


def a_docx(path):
    import docx

    document = docx.Document()
    document.add_paragraph("The palladium core degrades under sustained load.")
    document.add_paragraph("Keep a spare housing machined and ready.")
    table = document.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "part"
    table.cell(0, 1).text = "hours"
    table.cell(1, 0).text = "housing"
    table.cell(1, 1).text = "2"
    document.save(str(path))
    return path


def a_workbook(path, rows=5):
    from openpyxl import Workbook

    workbook = Workbook()
    sheet = workbook.active
    sheet.title = "readings"
    sheet.append(["hour", "output"])
    for hour in range(rows):
        sheet.append([hour, hour * 3])
    workbook.create_sheet("notes").append(["nothing yet"])
    workbook.save(str(path))
    return path


def a_deck(path, slides=("Reactor status", "Next steps & timing")):
    """A minimal .pptx: a zip of XML, which is all a deck is."""
    with zipfile.ZipFile(path, "w") as archive:
        archive.writestr("[Content_Types].xml", "<Types/>")
        for index, body in enumerate(slides, 1):
            archive.writestr(
                f"ppt/slides/slide{index}.xml",
                f"<p:sld><a:t>{body.replace('&', '&amp;')}</a:t></p:sld>")
    return path


# ---------------------------------------------------------------------------
# PDF
# ---------------------------------------------------------------------------


def test_a_pdf_comes_back_as_text(root, run):
    result = D.documents_extract(run, str(a_pdf(root / "log.pdf")))
    assert result.status == c.SUCCEEDED
    assert "Arc reactor" in result.output["text"]
    assert result.output["pages"] == 2


def test_a_page_range_reads_only_those_pages(root, run):
    a_pdf(root / "log.pdf", pages=("first page", "second page", "third page"))
    result = D.documents_extract(run, str(root / "log.pdf"), pages="2")
    assert "second page" in result.output["text"]
    assert "first page" not in result.output["text"]
    assert result.output["pages_read"] == [2]


def test_a_page_range_outside_the_document_reads_all_of_it(root, run):
    a_pdf(root / "log.pdf")
    result = D.documents_extract(run, str(root / "log.pdf"), pages="9-12")
    assert result.output["pages_read"] == [1, 2], "a bad range lost the document"


def test_a_scanned_pdf_says_it_needs_ocr_rather_than_reporting_nothing(root, run):
    """
    "0 characters" is true and useless. The file was read correctly; it has no
    text in it, and Friday has no OCR.
    """
    result = D.documents_extract(run, str(a_blank_pdf(root / "scan.pdf")))
    assert result.status == c.PARTIAL
    assert result.output["looks_scanned"] is True
    assert "OCR" in (result.error or "")


# ---------------------------------------------------------------------------
# Word, spreadsheets, decks, archives
# ---------------------------------------------------------------------------


def test_a_word_document_includes_its_tables(root, run):
    result = D.documents_extract(run, str(a_docx(root / "notes.docx")))
    assert "palladium core" in result.output["text"]
    assert "housing | 2" in result.output["text"], "the table was dropped"
    assert result.output["tables"] == 1


def test_a_spreadsheet_comes_back_as_csv(root, run):
    result = D.documents_extract(run, str(a_workbook(root / "book.xlsx")))
    assert result.output["text"].splitlines()[0] == "hour,output"
    assert result.output["sheet_read"] == "readings"
    assert result.output["sheets"] == ["readings", "notes"]


def test_a_named_sheet_is_read_and_a_missing_one_is_refused(root, run):
    a_workbook(root / "book.xlsx")
    named = D.documents_extract(run, str(root / "book.xlsx"), sheet="notes")
    assert "nothing yet" in named.output["text"]

    missing = D.documents_extract(run, str(root / "book.xlsx"), sheet="ghost")
    assert missing.status == c.FAILED
    assert "no sheet" in missing.error


def test_a_huge_spreadsheet_stops_and_says_so(root, run, monkeypatch):
    monkeypatch.setattr(D, "MAX_ROWS", 4)
    result = D.documents_extract(run, str(a_workbook(root / "big.xlsx", rows=50)))
    assert result.output["rows_read"] == 4
    assert result.output["truncated_rows"] is True, "a silent cap"


def test_a_slide_deck_comes_back_slide_by_slide(root, run):
    result = D.documents_extract(run, str(a_deck(root / "deck.pptx")))
    assert "[slide 1] Reactor status" in result.output["text"]
    assert "Next steps & timing" in result.output["text"], "entities not decoded"
    assert result.output["slides"] == 2


def test_an_archive_is_listed_not_unpacked(root, run):
    """
    Listing is a read; extracting is a write, and a zip that unpacks
    `../../../.ssh/id_rsa` is the oldest trick there is.
    """
    archive = root / "bundle.zip"
    with zipfile.ZipFile(archive, "w") as handle:
        handle.writestr("inner/one.txt", "hello")
        handle.writestr("../escape.txt", "nope")
    before = {p.name for p in root.iterdir()}

    result = D.documents_extract(run, str(archive))
    assert "inner/one.txt" in result.output["text"]
    assert result.output["entries"] == 2
    assert {p.name for p in root.iterdir()} == before, "it unpacked something"


# ---------------------------------------------------------------------------
# Inspect is the cheap question
# ---------------------------------------------------------------------------


def test_inspecting_a_pdf_counts_pages_without_the_text(root, run):
    result = D.documents_inspect(run, str(a_pdf(root / "log.pdf")))
    assert result.status == c.SUCCEEDED
    assert result.output["pages"] == 2
    assert "text" not in result.output


def test_inspecting_a_workbook_names_its_sheets(root, run):
    result = D.documents_inspect(run, str(a_workbook(root / "book.xlsx")))
    assert result.output["sheets"] == ["readings", "notes"]


# ---------------------------------------------------------------------------
# The boundary
# ---------------------------------------------------------------------------


def test_a_document_outside_the_roots_is_refused(root, run, tmp_path):
    outside = tmp_path.parent / "elsewhere.pdf"
    a_pdf(outside)
    result = D.documents_extract(run, str(outside))
    assert result.status == c.FAILED
    assert "outside the permitted roots" in result.error


def test_a_protected_file_is_refused_even_with_a_document_suffix(root, run):
    """
    A document reader pointed at a credentials file is an exfiltration
    primitive wearing a helpful name.
    """
    secret = root / ".env.pdf"
    a_pdf(secret)
    result = D.documents_extract(run, str(secret))
    assert result.status == c.FAILED
    assert "protected pattern" in result.error


def test_a_plain_text_file_is_sent_to_files_read(root, run):
    plain = root / "notes.txt"
    plain.write_text("just text", encoding="utf-8")
    result = D.documents_extract(run, str(plain))
    assert result.status == c.FAILED
    assert "files_read" in result.error, "it did not say where to go instead"


def test_a_corrupt_document_fails_as_a_result_not_an_exception(root, run):
    broken = root / "broken.pdf"
    broken.write_bytes(b"this is not a pdf at all")
    result = D.documents_extract(run, str(broken))
    assert result.status == c.FAILED
    assert "could not read" in result.error


def test_nothing_here_can_execute_anything(root, run):
    """
    The donor's `action="run"` did `subprocess.run(["python", str(path)])` -
    an arbitrary execution primitive reachable by naming a file. This keeps it
    out at the module level, because the next person to port a handler will
    not read the docstring first.

    Parsed, not grepped. The first version of this searched the source text
    for "subprocess" and failed on the module docstring, which mentions it
    while explaining that it is absent - the fifth time in this project that a
    source-matching test has caught prose ABOUT the behaviour instead of the
    behaviour. An import is an AST node; a sentence is not.
    """
    import ast
    import inspect

    tree = ast.parse(inspect.getsource(D))
    imported, builtins_called, methods_called = set(), set(), set()
    for node in ast.walk(tree):
        if isinstance(node, ast.Import):
            imported.update(alias.name.split(".")[0] for alias in node.names)
        elif isinstance(node, ast.ImportFrom) and node.module:
            imported.add(node.module.split(".")[0])
        elif isinstance(node, ast.Call):
            if isinstance(node.func, ast.Name):
                builtins_called.add(node.func.id)
            elif isinstance(node.func, ast.Attribute):
                methods_called.add(node.func.attr)

    assert not imported & {"subprocess", "pty", "ctypes"}, sorted(imported)
    # Bare names and attributes are checked separately, because `compile` on
    # its own is the builtin that turns text into code and `re.compile` is a
    # regex. Lumping them together fails on the second, which is how a guard
    # gets deleted for being noisy.
    assert not builtins_called & {"eval", "exec", "compile", "__import__"}, \
        sorted(builtins_called)
    assert not methods_called & {"system", "popen", "Popen", "spawnv",
                                 "spawnl", "execv", "check_output"}, \
        sorted(methods_called)


def test_files_read_points_at_the_document_reader_instead_of_stopping(root, run):
    """
    The collision, measured live: asked "what does the pdf at <path> say", the
    model reached for files_read - which is CORE, always visible, and costs no
    search - got "binary", and told the boss "I can't read its contents
    directly. It's a binary file." It never searched for a PDF capability,
    because it already had a plausible tool in front of it and that tool's
    answer was a dead end.

    Ranking documents_extract first offline does nothing when discovery is
    never run. A result that names its own successor does.
    """
    result = F.files_read(run, str(a_pdf(root / "log.pdf")))
    assert result.status == c.PARTIAL
    assert "documents_extract" in (result.error or "")
    assert result.output["try_instead"] == "documents_extract"


def test_a_binary_with_no_reader_still_just_says_binary(root, run):
    """No successor to name, so none is invented."""
    blob = root / "mystery.bin"
    blob.write_bytes(b"\x00\x01\x02\xff\xfe")
    result = F.files_read(run, str(blob))
    assert result.status == c.PARTIAL
    assert "documents_extract" not in (result.error or "")
    assert result.output["try_instead"] == ""


def test_a_document_read_is_auto_and_needs_no_approval(root, run):
    """It is a read, and it is inside the jail. Asking would be theatre."""
    guarded = P.PolicyEngine(autonomy=P.GUARDED)
    result = D.documents_extract(run, str(a_pdf(root / "log.pdf")),
                                 engine=guarded)
    assert result.status == c.SUCCEEDED
